"""Combined Log File Tools Application
Combines:
- Log File Extract Tool
- Log File Check Tool
- Log File Parse Tool
- Log File Export Tool
into a single tabbed interface.
"""

import os
import tkinter as tk
from tkinter import ttk
from tkinter import filedialog as fdlg
from tkinter import messagebox
from pathlib import Path
import re
import shutil
import pandas as pd
import openpyxl
from openpyxl.utils import get_column_letter
from shutil import copyfile
from docx import Document
from docx.shared import Pt


# ==============================================================================
# TAB 1: Log File Extract Tool
# ==============================================================================


def clone_structure_with_logs(source_folder: str, destination_folder: str) -> None:
    """
    Clones the folder structure from source_folder downwards into destination_folder,
    copying all .log files into their respective mirrored directories.
    """
    source_path = Path(source_folder)
    destination_path = Path(destination_folder)

    if not source_path.exists():
        raise FileNotFoundError(f"Source folder not found: {source_folder}")
    if not source_path.is_dir():
        raise NotADirectoryError(f"Source path is not a directory: {source_folder}")

    for current_dir in source_path.rglob("*"):
        if current_dir.is_dir():
            mirrored_dir = destination_path / current_dir.relative_to(source_path)
            mirrored_dir.mkdir(parents=True, exist_ok=True)
        elif current_dir.suffix == ".log":
            mirrored_dir = destination_path / current_dir.parent.relative_to(
                source_path
            )
            mirrored_dir.mkdir(parents=True, exist_ok=True)
            shutil.copy2(current_dir, mirrored_dir / current_dir.name)

    print(f"Done. Cloned structure and copied .log files to: {destination_path}")


class ExtractTool(ttk.Frame):
    """Log File Extract Tool — clones folder structure and copies .log files."""

    def __init__(self, parent):
        super().__init__(parent)
        self.selected_topfolder = ""
        self.selected_destination = ""
        self._build_ui()

    def _build_ui(self):
        layout_frame = ttk.Frame(master=self)
        layout_frame.grid(column=0, row=0, columnspan=1, rowspan=1, sticky=tk.W + tk.E)
        layout_frame.columnconfigure(0, weight=1)
        self.main_title_label = ttk.Label(
            master=layout_frame,
            text="Log File Extract Tool",
            relief=tk.FLAT,
            font=("Times New Roman", 18, "bold"),
        )
        self.main_title_label.grid(
            row=0, column=0, columnspan=2, rowspan=1, ipadx=5, pady=0
        )

        self.select_topfolder_button = ttk.Button(
            master=layout_frame,
            text="Select Top Folder",
            command=self.select_top_folder,
        )
        self.select_topfolder_button.grid(
            row=1, column=0, columnspan=1, rowspan=1, padx=5, pady=5, sticky=tk.W + tk.E
        )
        self.select_topfolder_label_frame = ttk.LabelFrame(
            master=layout_frame, text="Top Folder", height=2
        )
        self.select_topfolder_label = ttk.Label(
            master=self.select_topfolder_label_frame, text="", width=100
        )
        self.select_topfolder_label_frame.grid(
            row=1, column=1, columnspan=1, rowspan=1, padx=5, pady=5, sticky=tk.W + tk.E
        )
        self.select_topfolder_label.grid(
            row=0, column=0, columnspan=1, rowspan=1, padx=5, pady=5, sticky=tk.W + tk.E
        )

        self.select_destination_button = ttk.Button(
            master=layout_frame,
            text="Select Destination Folder",
            command=self.select_destination,
        )
        self.select_destination_button.grid(
            row=2, column=0, columnspan=1, rowspan=1, padx=5, pady=5, sticky=tk.W + tk.E
        )
        self.select_destination_label_frame = ttk.LabelFrame(
            master=layout_frame, text="Destination", height=2
        )
        self.select_destination_label = ttk.Label(
            master=self.select_destination_label_frame, text="", width=100
        )
        self.select_destination_label_frame.grid(
            row=2, column=1, columnspan=1, rowspan=1, padx=5, pady=5, sticky=tk.W + tk.E
        )
        self.select_destination_label.grid(
            row=0, column=0, columnspan=1, rowspan=1, padx=5, pady=5, sticky=tk.W + tk.E
        )

        self.clone_structurefiles_button = ttk.Button(
            master=layout_frame,
            text="Clone",
            command=self.clone_filestructure,
        )
        self.clone_structurefiles_button.grid(
            row=3, column=0, columnspan=1, rowspan=1, padx=5, pady=5, sticky=tk.W + tk.E
        )

    def select_top_folder(self):
        target_dir = str(fdlg.askdirectory())
        self.select_topfolder_label.configure(text=target_dir)
        self.selected_topfolder = target_dir

    def select_destination(self):
        target_dir = str(fdlg.askdirectory())
        self.select_destination_label.configure(text=target_dir)
        self.selected_destination = target_dir

    def clone_filestructure(self):
        try:
            clone_structure_with_logs(
                source_folder=self.selected_topfolder,
                destination_folder=self.selected_destination,
            )
            messagebox.showinfo("Done", "Folder structure cloned successfully.")
        except Exception as e:
            messagebox.showerror("Error", str(e))


# ==============================================================================
# TAB 2: Log File Check Tool
# ==============================================================================


class CheckTool(ttk.Frame):
    """Log File Check Tool — checks which log files are in/out of the archive."""

    def __init__(self, parent):
        super().__init__(parent)
        self.target_archive_file = ""
        self.target_dir = ""
        self.data_dict = {}
        self.sheet_name = "Sheet1"
        self.not_in_archive_paths = {}
        self._build_ui()

    def _build_ui(self):
        # Title Frame
        title_frame = ttk.Frame(master=self)
        title_frame.grid(column=0, row=0, columnspan=1, rowspan=1, sticky=tk.W + tk.E)
        title_frame.columnconfigure(0, weight=1)

        self.main_title_label = ttk.Label(
            master=title_frame,
            text="Log File Check Tool",
            relief=tk.FLAT,
            font=("Times New Roman", 18, "bold"),
        )
        self.main_title_label.grid(
            row=0,
            column=0,
            columnspan=4,
            rowspan=1,
            ipadx=5,
            pady=0,
        )
        self.version_title_label = ttk.Label(
            master=title_frame,
            text="v1.0.0",
            font=("Times New Roman", 8, "italic"),
        )
        self.version_title_label.grid(
            row=1, column=0, columnspan=4, rowspan=1, padx=0, pady=0, sticky=tk.N
        )

        sep_1 = ttk.Separator(master=self, orient="horizontal")
        sep_1.grid(row=1, column=0, sticky="nesw")

        # Start Frame
        start_frame = ttk.Frame(master=self)
        start_frame.grid(
            column=0, row=2, columnspan=1, rowspan=1, padx=5, pady=5, sticky=tk.W + tk.E
        )
        start_frame.columnconfigure(1, weight=1)

        self.select_archive_file_button = ttk.Button(
            master=start_frame,
            text="Select Archive File",
            command=self.select_archive_file,
        )
        self.select_archive_file_button.grid(
            row=0, column=0, columnspan=1, rowspan=1, padx=5, pady=5, sticky=tk.W + tk.E
        )
        self.select_archive_label_frame = ttk.LabelFrame(
            master=start_frame, text="Selected Archive File", height=2
        )
        self.select_archive_file_label = ttk.Label(
            master=self.select_archive_label_frame, text="", width=130, anchor=tk.W
        )
        self.select_archive_label_frame.grid(
            row=0, column=1, columnspan=1, rowspan=1, padx=5, pady=5, sticky=tk.W + tk.E
        )
        self.select_archive_file_label.grid(
            row=0,
            column=0,
            columnspan=1,
            rowspan=1,
            padx=5,
            pady=5,
            sticky=tk.N + tk.W + tk.E,
        )

        self.archivefilekeep_yesno = tk.IntVar()
        self.keep_archive_file_checkbutton = ttk.Checkbutton(
            master=start_frame,
            text="Keep when reset?",
            onvalue=1,
            offvalue=0,
            variable=self.archivefilekeep_yesno,
        )
        self.keep_archive_file_checkbutton.grid(
            row=0,
            column=6,
            columnspan=1,
            rowspan=1,
            padx=5,
            pady=5,
            sticky=tk.W + tk.N + tk.S,
        )
        self.archivefilekeep_yesno.set(value=0)

        self.select_folders_button = ttk.Button(
            master=start_frame,
            text="Select Data Folder(s)",
            command=self.select_folders,
        )
        self.select_folders_button.grid(
            row=1, column=0, columnspan=1, rowspan=1, padx=5, pady=5, sticky=tk.W + tk.E
        )
        self.select_folders_label_frame = ttk.LabelFrame(
            master=start_frame, text="Selected Folder(s)", height=2
        )
        self.select_folders_label = ttk.Label(
            master=self.select_folders_label_frame, text="", width=130
        )
        self.select_folders_label_frame.grid(
            row=1, column=1, columnspan=1, rowspan=1, padx=5, pady=5, sticky=tk.W + tk.E
        )
        self.select_folders_label.grid(
            row=0,
            column=0,
            columnspan=1,
            rowspan=1,
            padx=5,
            pady=5,
            sticky=tk.N + tk.W + tk.E,
        )

        self.one_or_many_frame = ttk.LabelFrame(
            master=start_frame, text="One or Many", height=2
        )
        self.one_or_many_frame.grid(
            column=6, row=1, columnspan=1, rowspan=1, padx=5, pady=5, sticky=tk.W + tk.E
        )
        self.one_or_many_var = tk.IntVar()
        self.one_folder = ttk.Radiobutton(
            master=self.one_or_many_frame,
            text="One Folder",
            variable=self.one_or_many_var,
            value=0,
        )
        self.many_folder = ttk.Radiobutton(
            master=self.one_or_many_frame,
            text="Many Folders",
            variable=self.one_or_many_var,
            value=1,
        )
        self.one_folder.grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        self.many_folder.grid(row=1, column=0, padx=5, pady=5, sticky=tk.W)
        self.one_or_many_var.set(1)

        self.check_files_button = ttk.Button(
            master=start_frame, text="Check for Log Files", command=self.check_log_files
        )
        self.check_files_button.grid(
            row=2, column=0, columnspan=1, rowspan=1, padx=5, pady=5, sticky=tk.E + tk.W
        )

        sep_2 = ttk.Separator(master=self, orient="horizontal")
        sep_2.grid(row=3, column=0, sticky="nesw")

        # Input Frame (listboxes)
        input_frame = ttk.Frame(master=self)
        input_frame.grid(
            column=0,
            row=4,
            columnspan=1,
            rowspan=1,
            padx=5,
            pady=5,
            sticky=tk.W + tk.E + tk.N + tk.S,
        )
        for col in range(5):
            input_frame.columnconfigure(col, weight=1)

        # Empty Folders
        ttk.Label(
            master=input_frame,
            text="Empty Folders",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=0, column=0, columnspan=2, padx=5, pady=0, sticky=tk.W + tk.S)
        self.empty_folders_textvar = tk.StringVar()
        self.empty_folders_display = tk.Listbox(
            master=input_frame, listvariable=self.empty_folders_textvar, width=35
        )
        self.empty_folders_display.grid(
            row=1, column=0, columnspan=2, padx=5, pady=5, sticky=tk.W + tk.N + tk.E
        )
        self.empty_folders_textvar.set(value="")

        # Log Files in Archive
        ttk.Label(
            master=input_frame,
            text="Log Files in Archive",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=0, column=2, columnspan=2, padx=5, pady=0, sticky=tk.W + tk.S)
        self.log_files_in_archive_textvar = tk.StringVar()
        self.log_files_in_archive_display = tk.Listbox(
            master=input_frame, listvariable=self.log_files_in_archive_textvar, width=35
        )
        self.log_files_in_archive_display.grid(
            row=1, column=2, columnspan=2, padx=5, pady=5, sticky=tk.W + tk.N + tk.E
        )
        self.log_files_in_archive_textvar.set(value="")

        # Log Files NOT in Archive
        ttk.Label(
            master=input_frame,
            text="Log Files NOT in Archive",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=0, column=4, columnspan=2, padx=5, pady=0, sticky=tk.W + tk.S)
        self.log_files_not_in_archive_textvar = tk.StringVar()
        self.log_files_not_in_archive_display = tk.Listbox(
            master=input_frame,
            listvariable=self.log_files_not_in_archive_textvar,
            width=35,
        )
        self.log_files_not_in_archive_display.grid(
            row=1, column=4, columnspan=2, padx=5, pady=5, sticky=tk.W + tk.N + tk.E
        )
        self.log_files_not_in_archive_textvar.set(value="")

        self.reduce_recheck_button = ttk.Button(
            master=input_frame,
            text="Reduce and Recheck",
            command=self.reduce_and_recheck,
        )
        self.reduce_recheck_button.grid(
            row=2, column=4, columnspan=2, padx=5, pady=5, sticky=tk.W + tk.E
        )

        # Folders with Errors
        ttk.Label(
            master=input_frame,
            text="Folders with Errors",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=0, column=6, columnspan=2, padx=5, pady=0, sticky=tk.W + tk.S)
        self.error_folders_textvar = tk.StringVar()
        self.error_folders_display = tk.Listbox(
            master=input_frame, listvariable=self.error_folders_textvar, width=35
        )
        self.error_folders_display.grid(
            row=1, column=6, columnspan=2, padx=5, pady=5, sticky=tk.W + tk.N + tk.E
        )
        self.error_folders_textvar.set(value="")

        # Output Frame
        output_frame = ttk.Frame(master=self)
        output_frame.grid(
            column=0, row=5, columnspan=1, rowspan=1, padx=5, pady=5, sticky=tk.E
        )

        self.export_button = ttk.Button(
            master=output_frame,
            text="Export NOT in Archive",
            command=self.export_not_in_archive,
        )
        self.export_button.grid(row=2, column=0, padx=5, pady=5, sticky=tk.E)

        self.reset_clear_button = ttk.Button(
            master=output_frame, text="Reset / Clear", command=self.reset_gui
        )
        self.reset_clear_button.grid(row=2, column=1, padx=5, pady=5, sticky=tk.E)

    def select_archive_file(self):
        selected = fdlg.askopenfilename()
        self.select_archive_file_label.configure(text=str(selected))
        self.target_archive_file = selected

    def select_folders(self):
        target_dir = str(fdlg.askdirectory())
        self.select_folders_label.configure(text=target_dir)
        self.target_dir = target_dir

    def check_log_files(self):
        log_extension = ".log"
        log_file_list = []
        empty_list = []
        in_archive_list = []
        not_in_archive_list = []
        error_list = []

        if self.target_archive_file == "":
            messagebox.showwarning("Warning", "An Archive File needs to be selected.")
            return
        if self.target_dir == "":
            messagebox.showwarning(
                "Warning", "A target directory needs to be selected."
            )
            return

        directory_path = Path(self.target_dir)

        def filter_tilde_files(files):
            tilde_pattern = re.compile(r"^(.+)~\d+$")
            clean_stems = {f.stem for f in files if not tilde_pattern.match(f.stem)}
            filtered = []
            for f in files:
                match = tilde_pattern.match(f.stem)
                if match:
                    if match.group(1) not in clean_stems:
                        filtered.append(f)
                else:
                    filtered.append(f)
            return filtered

        def collect_log_files(folder):
            collected = []
            try:
                contents = list(Path(folder).iterdir())
                if not contents:
                    empty_list.append(str(Path(folder).stem))
                    return collected
                for item in contents:
                    if item.is_dir():
                        collected.extend(collect_log_files(item))
                    elif item.suffix == log_extension:
                        collected.append(item)
            except Exception as e:
                error_list.append(f"{Path(folder).stem} ({e})")
            return collected

        if int(self.one_or_many_var.get()) == 0:
            raw_files = collect_log_files(directory_path)
            log_file_list = filter_tilde_files(raw_files)
        elif int(self.one_or_many_var.get()) == 1:
            sub_dir_list = [x for x in directory_path.iterdir() if x.is_dir()]
            if not sub_dir_list:
                messagebox.showwarning("Warning", "No subdirectories found.")
                return
            for each_sub_dir in sub_dir_list:
                raw_files = collect_log_files(each_sub_dir)
                log_file_list.extend(filter_tilde_files(raw_files))

        archive_file = Path(self.target_archive_file)
        wb = pd.read_excel(archive_file, sheet_name=self.sheet_name)
        for each_log_file in log_file_list:
            if each_log_file.stem in wb.Filename.values:
                in_archive_list.append(each_log_file.stem)
            else:
                not_in_archive_list.append(each_log_file.stem)
                self.not_in_archive_paths[each_log_file.stem] = str(
                    each_log_file.resolve()
                )

        for each_item in empty_list:
            self.empty_folders_display.insert(tk.END, each_item)
        for each_item in in_archive_list:
            self.log_files_in_archive_display.insert(tk.END, each_item)
        for each_item in not_in_archive_list:
            self.log_files_not_in_archive_display.insert(tk.END, each_item)
        for each_item in error_list:
            self.error_folders_display.insert(tk.END, each_item)

    def reduce_and_recheck(self):
        tilde_pattern = re.compile(r"~\d+$")
        all_items = list(self.log_files_not_in_archive_display.get(0, tk.END)) + list(
            self.log_files_in_archive_display.get(0, tk.END)
        )
        cleaned = set()
        for item in all_items:
            cleaned.add(re.sub(tilde_pattern, "", item))

        self.log_files_in_archive_display.delete(0, tk.END)
        self.log_files_not_in_archive_display.delete(0, tk.END)

        try:
            archive_file = Path(self.target_archive_file)
            wb = pd.read_excel(archive_file, sheet_name=self.sheet_name)
            for stem in sorted(cleaned):
                if stem in wb.Filename.values:
                    self.log_files_in_archive_display.insert(tk.END, stem)
                else:
                    self.log_files_not_in_archive_display.insert(tk.END, stem)
                    original = next(
                        (
                            k
                            for k in self.not_in_archive_paths
                            if re.sub(tilde_pattern, "", k) == stem
                        ),
                        None,
                    )
                    if original:
                        self.not_in_archive_paths[stem] = re.sub(
                            tilde_pattern, "", self.not_in_archive_paths[original]
                        )
        except Exception as e:
            self.error_folders_display.insert(tk.END, f"Recheck error: {e}")

    def export_not_in_archive(self):
        items = list(self.log_files_not_in_archive_display.get(0, tk.END))
        if not items:
            messagebox.showinfo("Info", "No items to export.")
            return
        save_path = fdlg.asksaveasfilename(
            defaultextension=".txt",
            initialfile="Files_to_Archive.txt",
            filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")],
        )
        if not save_path:
            return
        try:
            with open(save_path, "w", encoding="utf-8") as f:
                for item in items:
                    path = self.not_in_archive_paths.get(item, item)
                    f.write(path + "\n")
            messagebox.showinfo("Done", f"Exported {len(items)} items to {save_path}")
        except Exception as e:
            self.error_folders_display.insert(tk.END, f"Export error: {e}")

    def reset_gui(self):
        if int(self.archivefilekeep_yesno.get()) == 0:
            self.select_archive_file_label.configure(text="")
        self.select_folders_label.configure(text="")
        self.empty_folders_display.delete(0, tk.END)
        self.log_files_in_archive_display.delete(0, tk.END)
        self.log_files_not_in_archive_display.delete(0, tk.END)
        self.error_folders_display.delete(0, tk.END)
        self.not_in_archive_paths = {}


# ==============================================================================
# TAB 3: Log File Parse Tool
# ==============================================================================


class ParseTool(ttk.Frame):
    """Log File Parse Tool — parses log files and publishes data to an archive."""

    def __init__(self, parent):
        super().__init__(parent)
        self.selected_log_source_file = ""
        self.target_archive_file = ""
        self.target_dir = ""
        self.log_file_text = ""
        self.data_dict = {}
        self.batch_file_list = []
        self.batch_current_index = -1
        self.batch_selected_index = -1
        self._build_ui()

    def _build_ui(self):
        # Title Frame
        title_frame = ttk.Frame(master=self)
        title_frame.grid(column=0, row=0, columnspan=1, rowspan=1, sticky=tk.W + tk.E)
        title_frame.columnconfigure(0, weight=1)

        self.main_title_label = ttk.Label(
            master=title_frame,
            text="Log File Parse Tool",
            relief=tk.FLAT,
            font=("Times New Roman", 18, "bold"),
        )
        self.main_title_label.grid(
            row=0,
            column=0,
            columnspan=4,
            ipadx=5,
            pady=0,
        )
        self.version_title_label = ttk.Label(
            master=title_frame, text="v2.1.0", font=("Times New Roman", 8, "italic")
        )
        self.version_title_label.grid(
            row=1, column=0, columnspan=4, padx=0, pady=0, sticky=tk.N
        )
        self.help_button = ttk.Button(
            master=title_frame, text="Help", command=self.open_help_box
        )
        self.help_button.grid(row=1, column=4, padx=5, pady=5, sticky=tk.E)

        sep_1 = ttk.Separator(master=self, orient="horizontal")
        sep_1.grid(row=1, column=0, sticky="nesw")

        # Start Frame
        start_frame = ttk.Frame(master=self)
        start_frame.grid(
            column=0, row=2, columnspan=1, rowspan=1, padx=5, pady=5, sticky=tk.W + tk.E
        )
        start_frame.columnconfigure(1, weight=1)

        # Destination File
        self.select_destination_file_button = ttk.Button(
            master=start_frame,
            text="Select Destination File",
            command=self.select_destination_file,
        )
        self.select_destination_file_button.grid(
            row=0, column=0, padx=5, pady=5, sticky=tk.W
        )
        self.select_destination_file_label_frame = ttk.LabelFrame(
            master=start_frame, text="Destination File Name", height=2
        )
        self.select_destination_file_label = ttk.Label(
            master=self.select_destination_file_label_frame,
            text="",
            width=130,
            anchor=tk.W,
        )
        self.select_destination_file_label_frame.grid(
            row=0, column=1, padx=5, pady=5, sticky=tk.W + tk.E
        )
        self.select_destination_file_label.grid(
            row=0, column=0, padx=5, pady=5, sticky=tk.N + tk.W + tk.E
        )
        self.indicator_light_destination_file = tk.Canvas(
            master=start_frame, bg="white", height=50, width=100
        )
        self.indicator_light_destination_file.grid(column=5, row=0, sticky=tk.E)
        self.indicator_light_destination_file.create_text(
            30, 25, text="Selection\nIndicator"
        )
        self.destination_file_light = self.indicator_light_destination_file.create_oval(
            75, 20, 90, 35, fill="red"
        )
        self.destinationfilekeep_yesno = tk.IntVar()
        ttk.Checkbutton(
            master=start_frame,
            text="Keep when reset?",
            onvalue=1,
            offvalue=0,
            variable=self.destinationfilekeep_yesno,
        ).grid(row=0, column=6, padx=5, pady=5, sticky=tk.W + tk.N + tk.S)
        self.destinationfilekeep_yesno.set(0)

        # Log Source File
        self.select_log_file_button = ttk.Button(
            master=start_frame,
            text="Select Log Source File",
            command=self.select_log_source_file,
        )
        self.select_log_file_button.grid(row=1, column=0, padx=5, pady=5, sticky=tk.W)
        self.select_log_file_label_frame = ttk.LabelFrame(
            master=start_frame, text="Log File Name", height=2
        )
        self.select_log_file_label = ttk.Label(
            master=self.select_log_file_label_frame, text="", width=130
        )
        self.select_log_file_label_frame.grid(
            row=1, column=1, padx=5, pady=5, sticky=tk.W + tk.E
        )
        self.select_log_file_label.grid(
            row=0, column=0, padx=5, pady=5, sticky=tk.N + tk.W + tk.E
        )
        self.import_data_button = ttk.Button(
            master=start_frame, text="Parse Data", command=self.parse_data
        )
        self.import_data_button.grid(row=1, column=6, padx=5, pady=5)
        self.indicator_light_importdata = tk.Canvas(
            master=start_frame, bg="white", height=50, width=100
        )
        self.indicator_light_importdata.grid(column=5, row=1, sticky=tk.E)
        self.indicator_light_importdata.create_text(30, 25, text="Selection\nIndicator")
        self.import_data_light = self.indicator_light_importdata.create_oval(
            75, 20, 90, 35, fill="red"
        )

        # Log File Destination
        self.shift_log_file_button = ttk.Button(
            master=start_frame,
            text="Select Log File Destination",
            command=self.shift_log_file,
        )
        self.shift_log_file_button.grid(row=2, column=0, padx=5, pady=5, sticky=tk.E)
        self.shift_log_file_label_frame = ttk.LabelFrame(
            master=start_frame, text="Log File Destination", height=2
        )
        self.shift_log_file_label = ttk.Label(
            master=self.shift_log_file_label_frame, text="", width=130
        )
        self.shift_log_file_label_frame.grid(
            row=2, column=1, padx=5, pady=5, sticky=tk.W + tk.E
        )
        self.shift_log_file_label.grid(
            row=0, column=0, padx=5, pady=5, sticky=tk.W + tk.E
        )
        self.indicator_light_shift_logfile = tk.Canvas(
            master=start_frame, bg="white", height=50, width=100
        )
        self.indicator_light_shift_logfile.grid(row=2, column=5, sticky=tk.E)
        self.indicator_light_shift_logfile.create_text(
            30, 25, text="Selection\nIndicator"
        )
        self.shift_log_file_light = self.indicator_light_shift_logfile.create_oval(
            75, 20, 90, 35, fill="red"
        )
        self.destinationdirectorykeep_yesno = tk.IntVar()
        ttk.Checkbutton(
            master=start_frame,
            text="Keep when reset?",
            onvalue=1,
            offvalue=0,
            variable=self.destinationdirectorykeep_yesno,
        ).grid(row=2, column=6, padx=5, pady=5, sticky=tk.W + tk.N + tk.S)
        self.destinationdirectorykeep_yesno.set(0)

        # Batch File
        self.select_batch_file_button = ttk.Button(
            master=start_frame, text="Select Batch File", command=self.select_batch_file
        )
        self.select_batch_file_button.grid(row=3, column=0, padx=5, pady=5, sticky=tk.W)
        self.select_batch_file_label_frame = ttk.LabelFrame(
            master=start_frame, text="Batch File", height=2
        )
        self.select_batch_file_label = ttk.Label(
            master=self.select_batch_file_label_frame, text="", width=130
        )
        self.select_batch_file_label_frame.grid(
            row=3, column=1, padx=5, pady=5, sticky=tk.W + tk.E
        )
        self.select_batch_file_label.grid(
            row=0, column=0, padx=5, pady=5, sticky=tk.N + tk.W + tk.E
        )

        ttk.Label(
            master=start_frame,
            text="Batch Queue:",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=4, column=0, padx=5, pady=(5, 0), sticky=tk.W)
        self.batch_listbox = tk.Listbox(master=start_frame, width=80, height=6)
        self.batch_listbox.grid(
            row=5, column=0, columnspan=6, padx=5, pady=(0, 5), sticky=tk.W + tk.E
        )
        self.batch_listbox.bind("<<ListboxSelect>>", self.on_batch_select)

        sep_2 = ttk.Separator(master=self, orient="horizontal")
        sep_2.grid(row=3, column=0, sticky="nesw")

        # Input Frame
        input_frame = ttk.Frame(master=self)
        input_frame.grid(
            column=0,
            row=4,
            columnspan=1,
            rowspan=1,
            padx=5,
            pady=5,
            sticky=tk.W + tk.E + tk.N + tk.S,
        )
        for col in range(5):
            input_frame.columnconfigure(col, weight=1)

        # --- Row 0-1: Proposal, Name, Date, Position, Operator ---
        ttk.Label(
            master=input_frame,
            text="Proposal #:",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=0, column=0, padx=5, pady=0, sticky=tk.W + tk.S)
        self.proposalnumber_textvar = tk.StringVar(value="Enter Here")
        ttk.Entry(
            master=input_frame, textvariable=self.proposalnumber_textvar, width=35
        ).grid(row=1, column=0, padx=5, pady=5, sticky=tk.W + tk.N)

        ttk.Label(
            master=input_frame,
            text="PI and Collaborators",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=0, column=1, padx=5, pady=0, sticky=tk.W + tk.S)
        self.name_textvar = tk.StringVar(value="Enter Here")
        ttk.Entry(master=input_frame, textvariable=self.name_textvar, width=35).grid(
            row=1, column=1, padx=5, pady=5, sticky=tk.W + tk.N
        )

        ttk.Label(
            master=input_frame,
            text="Date:",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=0, column=2, padx=5, pady=0, sticky=tk.W + tk.S)
        self.date_textvar = tk.StringVar(value="Known")
        self.date_input = ttk.Entry(
            master=input_frame, textvariable=self.date_textvar, width=35
        )
        self.date_input.grid(row=1, column=2, padx=5, pady=5, sticky=tk.W + tk.N)

        ttk.Label(
            master=input_frame,
            text="Position:",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=0, column=3, padx=5, pady=0, sticky=tk.W)
        checkbox_frame = tk.Frame(master=input_frame)
        checkbox_frame.grid(row=1, column=3, padx=2, pady=5, sticky=tk.W + tk.E)
        self.facultypos_intvar = tk.IntVar()
        self.staffpos_intvar = tk.IntVar()
        self.studentpos_intvar = tk.IntVar()
        self.otherpos_intvar = tk.IntVar()
        ttk.Checkbutton(
            master=checkbox_frame,
            text="Faculty",
            onvalue=1,
            offvalue=0,
            variable=self.facultypos_intvar,
        ).grid(row=0, column=0, pady=2, sticky="w")
        ttk.Checkbutton(
            master=checkbox_frame,
            text="Staff",
            onvalue=1,
            offvalue=0,
            variable=self.staffpos_intvar,
        ).grid(row=1, column=0, pady=2, sticky="w")
        ttk.Checkbutton(
            master=checkbox_frame,
            text="Student",
            onvalue=1,
            offvalue=0,
            variable=self.studentpos_intvar,
        ).grid(row=0, column=1, pady=2, sticky="w")
        ttk.Checkbutton(
            master=checkbox_frame,
            text="Other",
            onvalue=1,
            offvalue=0,
            variable=self.otherpos_intvar,
        ).grid(row=1, column=1, pady=2, sticky="w")

        ttk.Label(
            master=input_frame,
            text="Operator (Who did the actual scan):",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=0, column=4, padx=5, pady=0, sticky=tk.W)
        self.operator_textvar = tk.StringVar(value="Select Here")
        ttk.Entry(
            master=input_frame, textvariable=self.operator_textvar, width=35
        ).grid(row=1, column=4, padx=5, pady=5, sticky=tk.W)

        # --- Row 2-3: Internal/External, Filename, Voltage, Current, Resolution ---
        ttk.Label(
            master=input_frame,
            text="Internal/External/Both:",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=2, column=0, padx=5, pady=0, sticky=tk.W)
        self.nyitorext_textvar = tk.StringVar(value="Enter Here")
        ttk.Combobox(
            master=input_frame,
            textvariable=self.nyitorext_textvar,
            width=35,
            values=("Internal", "External", "Both"),
        ).grid(row=3, column=0, padx=5, pady=5, sticky=tk.W)

        ttk.Label(
            master=input_frame,
            text="Filename:",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=2, column=1, padx=5, pady=0, sticky=tk.W + tk.E)
        self.filename_textvar = tk.StringVar(value="Auto Fill")
        self.filename_input = ttk.Entry(
            master=input_frame, textvariable=self.filename_textvar, width=35
        )
        self.filename_input.grid(row=3, column=1, padx=5, pady=5, sticky=tk.W)

        ttk.Label(
            master=input_frame,
            text="Voltage:",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=2, column=2, padx=5, pady=0, sticky=tk.W)
        self.voltage_textvar = tk.StringVar(value="Auto Fill")
        self.voltage_input = ttk.Entry(
            master=input_frame, textvariable=self.voltage_textvar, width=20
        )
        self.voltage_input.grid(row=3, column=2, padx=5, pady=5, sticky=tk.W)

        ttk.Label(
            master=input_frame,
            text="Current:",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=2, column=3, padx=5, pady=0, sticky=tk.W)
        self.current_textvar = tk.StringVar(value="Auto Fill")
        self.current_input = ttk.Entry(
            master=input_frame, textvariable=self.current_textvar, width=20
        )
        self.current_input.grid(row=3, column=3, padx=5, pady=5, sticky=tk.W)

        ttk.Label(
            master=input_frame,
            text="Resolution:",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=2, column=4, padx=5, pady=0, sticky=tk.W)
        self.resolution_textvar = tk.StringVar(value="Select Here")
        self.resolution_combobox = ttk.Combobox(
            master=input_frame,
            textvariable=self.resolution_textvar,
            values=("0.5K", "1K", "2K"),
        )
        self.resolution_combobox.grid(
            row=3, column=4, padx=5, pady=5, sticky=tk.W + tk.N
        )

        # --- Row 4-5: Filter, Exposure, Pixel Size, Image Format, Type of Scan ---
        ttk.Label(
            master=input_frame,
            text="Filter:",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=4, column=0, padx=5, pady=0, sticky=tk.W)
        self.filter_textvar = tk.StringVar(value="Enter or Select Here")
        self.filter_combobox = ttk.Combobox(
            master=input_frame,
            textvariable=self.filter_textvar,
            values=(
                "None",
                "Br 0.25mm",
                "Al 1.0mm",
                "User Filter: Al 0.5mm",
                "User Filter: ",
            ),
        )
        self.filter_combobox.grid(row=5, column=0, padx=5, pady=5, sticky=tk.W)

        ttk.Label(
            master=input_frame,
            text="Exposure:",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=4, column=1, padx=5, pady=0, sticky=tk.W)
        self.exposure_textvar = tk.StringVar(value="Auto Fill")
        self.exposure_input = ttk.Entry(
            master=input_frame, textvariable=self.exposure_textvar, width=20
        )
        self.exposure_input.grid(row=5, column=1, padx=5, pady=5, sticky=tk.W)

        ttk.Label(
            master=input_frame,
            text="Pixel Size:",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=4, column=2, padx=5, pady=0, sticky=tk.W)
        self.pixelsize_textvar = tk.StringVar(value="Auto Fill")
        self.pixel_size_input = ttk.Entry(
            master=input_frame, textvariable=self.pixelsize_textvar, width=20
        )
        self.pixel_size_input.grid(row=5, column=2, padx=5, pady=5, sticky=tk.W)

        ttk.Label(
            master=input_frame,
            text="Image Format:",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=4, column=3, padx=5, pady=0, sticky=tk.W)
        self.imageformat_textvar = tk.StringVar(value="Auto Fill")
        self.image_format_input = ttk.Entry(
            master=input_frame, textvariable=self.imageformat_textvar, width=20
        )
        self.image_format_input.grid(row=5, column=3, padx=5, pady=5, sticky=tk.W)

        ttk.Label(
            master=input_frame,
            text="Type of Scan:",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=4, column=4, padx=5, pady=0, sticky=tk.W)
        scantype_frame = tk.Frame(master=input_frame)
        scantype_frame.grid(row=5, column=4, padx=2, pady=5, sticky=tk.W + tk.E)
        self.typeofscan_textvar = tk.StringVar(value="Enter Here")
        self.type_of_scan_input = ttk.Entry(
            master=scantype_frame, textvariable=self.typeofscan_textvar, width=35
        )
        self.type_of_scan_input.grid(row=0, column=0, padx=0, pady=0, sticky=tk.W)
        self.offset_yesno = tk.IntVar(value=0)
        ttk.Checkbutton(
            master=scantype_frame,
            text="Offset?",
            onvalue=1,
            offvalue=0,
            variable=self.offset_yesno,
        ).grid(row=0, column=1, padx=5, pady=5, sticky=tk.E)

        # --- Row 6-7: Rotation Step, Frames, Random Movement, 360, Scan Duration ---
        ttk.Label(
            master=input_frame,
            text="Rotation Step:",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=6, column=0, padx=5, pady=0, sticky=tk.W)
        self.rotationstep_textvar = tk.StringVar(value="Auto Fill")
        self.rotation_step_input = ttk.Entry(
            master=input_frame, textvariable=self.rotationstep_textvar, width=35
        )
        self.rotation_step_input.grid(row=7, column=0, padx=5, pady=5, sticky=tk.W)

        ttk.Label(
            master=input_frame,
            text="Frames:",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=6, column=1, padx=5, pady=0, sticky=tk.W)
        self.frames_textvar = tk.StringVar(value="Auto Fill")
        self.frames_input = ttk.Entry(
            master=input_frame, textvariable=self.frames_textvar, width=20
        )
        self.frames_input.grid(row=7, column=1, padx=5, pady=5, sticky=tk.W)

        ttk.Label(
            master=input_frame,
            text="Random Movement:",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=6, column=2, padx=5, pady=0, sticky=tk.W)
        self.randommovement_textvar = tk.StringVar(value="Auto Fill")
        self.random_movement_input = ttk.Entry(
            master=input_frame, textvariable=self.randommovement_textvar
        )
        self.random_movement_input.grid(row=7, column=2, padx=5, pady=5, sticky=tk.W)

        self.rotate360_yesno = tk.IntVar(value=0)
        ttk.Checkbutton(
            master=input_frame,
            text="Rotate 360?",
            onvalue=1,
            offvalue=0,
            variable=self.rotate360_yesno,
        ).grid(row=7, column=3, padx=5, pady=5, sticky=tk.W)

        ttk.Label(
            master=input_frame,
            text="Scan Duration:",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=6, column=4, padx=5, pady=0, sticky=tk.W)
        self.scanduration_textvar = tk.StringVar(value="Auto Fill")
        self.scan_duration_input = ttk.Entry(
            master=input_frame, textvariable=self.scanduration_textvar, width=25
        )
        self.scan_duration_input.grid(row=7, column=4, padx=5, pady=5, sticky=tk.W)

        # --- Row 8-9: Comments, checkboxes ---
        ttk.Label(
            master=input_frame,
            text="Comments:",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=8, column=0, padx=5, pady=0, sticky=tk.W)
        self.comments_input = tk.Text(master=input_frame, height=4, width=72)
        self.comments_input.grid(
            row=9, column=0, columnspan=3, padx=5, pady=5, sticky=tk.W
        )
        self.comments_input.insert(tk.END, "Enter Here")

        self.logfilearchive_yesno = tk.IntVar(value=0)
        ttk.Checkbutton(
            master=input_frame,
            text="Log File Archived?",
            onvalue=1,
            offvalue=0,
            variable=self.logfilearchive_yesno,
        ).grid(row=9, column=3, padx=5, pady=5, sticky=tk.W + tk.N)
        self.visstorage_yesno = tk.IntVar(value=0)
        ttk.Checkbutton(
            master=input_frame,
            text="Added to Viscenter Storage?",
            onvalue=1,
            offvalue=0,
            variable=self.visstorage_yesno,
        ).grid(row=9, column=4, padx=5, pady=5, sticky=tk.W + tk.N)

        sep_3 = ttk.Separator(master=self, orient="horizontal")
        sep_3.grid(row=5, column=0, sticky="nesw")

        # Output Frame
        output_frame = ttk.Frame(master=self)
        output_frame.grid(column=0, row=6, padx=5, pady=5, sticky=tk.E)

        ttk.Button(
            master=output_frame, text="Publish Info", command=self.publish_info
        ).grid(row=0, column=1, padx=5, pady=5, sticky=tk.E)
        ttk.Button(
            master=output_frame, text="Reset / Clear", command=self.reset_clear
        ).grid(row=2, column=1, padx=5, pady=5, sticky=tk.E)

    # --- Methods ---

    def open_help_box(self):
        hb = tk.Toplevel(self)
        hb.geometry("710x480")
        hb.title("Help Box Window")
        help_notebook = ttk.Notebook(master=hb)
        frames = [ttk.Frame(master=help_notebook) for _ in range(6)]
        help_notebook.grid(row=0, column=0, padx=5, pady=5)
        for f in frames:
            f.grid(row=0, column=0, padx=5, pady=5)
        tabs = [
            "General Info",
            "Destination File",
            "Log File",
            "Parsed Information",
            "Publishing Data",
            "Other",
        ]
        for f, t in zip(frames, tabs):
            help_notebook.add(f, text=t)

        help_texts = [
            """
            This is the Log File Parser Tool
            Designed and Created by Dan Gibbons
            Property of NYIT

            The purpose of this app is to make it easy to transfer log data from log files generated after use of the
            MicroCT machine to the designated archive file to make further analysis and visualization possible.

            For questions about this application not addressed in this help section or
            too complicated or underexplained, please contact dgibbo03@nyit.edu.

            For questions about data usage and publication queries, please contact khurdle@nyit.edu.
            """,
            """
            The first step of the process is selecting the Archive Target File (.xlsx)
            file where which the data extracted will end up being published. This can be selected by
            clicking the 'Select Destination File' button at the top left of the GUI.

            After selecting the archive file, the text display next to the button should display the
            global pathway to the now selected document. The indicator light to the right will now turn
            green, indicating selection has been carried out.

            NOTE: There is a checkbox to the right that says 'Keep When Reset?'.
            This may be selected when processing numerous log files when the target archive file is
            to remain the same for all of the log files.
            """,
            """
            The second step is selecting the log file from which you will be extracting the usage information.

            Click the button labeled 'Select Log Source File'. Once selected, its global pathway will be
            displayed to the right of the button and the indicator light will turn green.

            Only after the file is selected can you click the 'Parse Data' button.

            After parsing, the appropriate fields will populate with data from the log file selected.
            """,
            """
            The third step is filling in the missing info.

            After the proper data has been parsed, entry fields with 'Enter Here' or checkboxes
            will need to be filled in manually.

            Failure to fill an input field will result in the recorded data displaying 'Enter Here', False, or nothing.
            """,
            """
            After all fields have been filled out, clicking the 'Publish Info' button will publish the
            data to the indicated destination archive file.

            If the 'Log File Destination' section was filled out before 'Publish Info' was pressed,
            the source log file will be moved to the designated directory after data publishing is complete.
            """,
            """
            Field reference:
            - Proposal # -> established beforehand, known to the users before scanning
            - PI and Collaborators -> names of each person associated with the project
            - Date -> date of the scan
            - Position -> position(s) of the users
            - Operator -> person who operated the machine
            - Internal/External/Both -> NYIT affiliated, External, or Both
            - Filename -> name of the log file, not including extension
            - Voltage, Current, Resolution, Filter, Exposure, Pixel Size -> scan parameters
            - Image Format -> image format output by the scan
            - Type of Scan -> Regular, Oversized, Segment #, etc.
            - Rotation Step, Frames, Random Movement, Rotate 360 -> scan motion parameters
            - Scan Duration -> duration of the scan
            - Comments -> any additional notes
            - Log File Archived -> if the log file was/will be archived
            - Added to Viscenter Storage -> if data was added to Viscenter Storage
            """,
        ]
        for frame, text in zip(frames, help_texts):
            ttk.Label(master=frame, text=text).grid(row=0, column=0, padx=5, pady=5)

    def select_destination_file(self):
        selected = fdlg.askopenfilename()
        self.select_destination_file_label.configure(text=str(selected))
        if str(selected) != "":
            self.indicator_light_destination_file.itemconfig(
                self.destination_file_light, fill="green"
            )
            self.target_archive_file = selected

    def select_log_source_file(self):
        self.selected_log_source_file = fdlg.askopenfilename()
        self.select_log_file_label.configure(text=str(self.selected_log_source_file))
        if str(self.selected_log_source_file) != "":
            self.indicator_light_importdata.itemconfig(
                self.import_data_light, fill="yellow"
            )

    def parse_data(self):
        with open(str(self.selected_log_source_file), "r", encoding="utf-8") as f:
            self.log_file_text = f.read()

        month_dict = {
            "Jan": 1,
            "Feb": 2,
            "Mar": 3,
            "Apr": 4,
            "May": 5,
            "Jun": 6,
            "Jul": 7,
            "Aug": 8,
            "Sep": 9,
            "Oct": 10,
            "Nov": 11,
            "Dec": 12,
        }
        patterns = {
            "date": re.compile(
                r"Study Date and Time\s*=\s*(\w+)\s*(\d\d),\s*(\d\d\d\d).*", re.I
            ),
            "filename": re.compile(r"Filename Prefix\s*=\s*(.*)", re.I),
            "voltage": re.compile(r"Source Voltage \(kV\)\s*=\s*(\d+)", re.I),
            "current": re.compile(r"Source Current \(uA\)\s*=\s*(\d+)", re.I),
            "resolution": re.compile(r"Number of Rows\s*=\s*(\d+)", re.I),
            "exposure": re.compile(r"Exposure \(ms\)\s*=\s*(\d+)", re.I),
            "pixelsize": re.compile(r"Image Pixel Size\s*\(um\)\s*=(\d+\.\d+)", re.I),
            "filter": re.compile(r"Filter\s*=\s*(.*)"),
            "imageformat": re.compile(r"Image Format\s*=\s*(\w+)", re.I),
            "typeofscan": re.compile(r"Number of connected scans\s*=\s*(\d+)", re.I),
            "offset": re.compile(r"Camera Offset\s*=\s*(ON|OFF)", re.I),
            "rotationstep": re.compile(
                r"Rotation Step\s*\(deg\)\s*=\s*(\d+\.\d+)", re.I
            ),
            "frames": re.compile(r"Frame Averaging\s*=(ON|OFF)\s*\((\d+)\)", re.I),
            "randommovement": re.compile(
                r"Random Movement\s*=(ON|OFF)\s*\((\d+)\)", re.I
            ),
            "rotate360": re.compile(r"Use 360 Rotation\s*=\s*(\w+)", re.I),
            "scanduration": re.compile(
                r"Scan duration\s*=\s*(\d{2}:\d{2}:\d{2})", re.I
            ),
        }

        finds = {k: p.search(self.log_file_text) for k, p in patterns.items()}

        try:
            f = finds["date"]
            if f:
                self.date_textvar.set(
                    f"{month_dict[f.group(1)]}/{f.group(2)}/{f.group(3)}"
                )
        except Exception:
            print("Date not found.")
        try:
            if finds["filename"]:
                self.filename_textvar.set(finds["filename"].group(1))
        except Exception:
            print("Filename not found.")
        try:
            if finds["voltage"]:
                self.voltage_textvar.set(finds["voltage"].group(1))
        except Exception:
            print("Voltage not found.")
        try:
            if finds["current"]:
                self.current_textvar.set(finds["current"].group(1))
        except Exception:
            print("Current not found.")
        try:
            f = finds["resolution"]
            if f:
                res_map = {2240: "2K", 1120: "1K", 560: "0.5K"}
                self.resolution_textvar.set(res_map.get(int(f.group(1)), "None"))
        except Exception:
            print("Resolution not found.")
        try:
            f = finds["filter"]
            if f:
                filter_map = {
                    "Al 1.0mm": "Al 1.0mm",
                    "brass 0.25mm": "Br 0.25mm",
                    "No Filter": "None",
                    "User filter": "User Filter: Al 0.5mm",
                }
                self.filter_textvar.set(filter_map.get(f.group(1), "User Filter: "))
        except Exception:
            print("Filter not found.")
        try:
            if finds["exposure"]:
                self.exposure_textvar.set(finds["exposure"].group(1))
        except Exception:
            print("Exposure not found.")
        try:
            if finds["pixelsize"]:
                self.pixelsize_textvar.set(finds["pixelsize"].group(1))
        except Exception:
            print("Pixel size not found.")
        try:
            if finds["imageformat"]:
                self.imageformat_textvar.set(finds["imageformat"].group(1))
        except Exception:
            print("Image format not found.")
        try:
            f = finds["typeofscan"]
            if f:
                t = int(f.group(1))
                self.typeofscan_textvar.set(
                    "Regular" if t == 1 else f"Oversize ({t})" if t > 1 else "N/a"
                )
        except Exception:
            print("Type of scan not found.")
        try:
            f = finds["offset"]
            if f:
                self.offset_yesno.set(1 if f.group(1) == "ON" else 0)
        except Exception:
            print("Offset not found.")
        try:
            if finds["rotationstep"]:
                self.rotationstep_textvar.set(finds["rotationstep"].group(1))
        except Exception:
            print("Rotation step not found.")
        try:
            f = finds["frames"]
            if f:
                self.frames_textvar.set("OFF" if f.group(1) == "OFF" else f.group(2))
        except Exception:
            print("Frames not found.")
        try:
            f = finds["randommovement"]
            if f:
                self.randommovement_textvar.set(
                    "OFF" if f.group(1) == "OFF" else f.group(2)
                )
        except Exception:
            print("Random movement not found.")
        try:
            f = finds["rotate360"]
            if f:
                self.rotate360_yesno.set(1 if f.group(1).upper() == "YES" else 0)
        except Exception:
            print("360 rotation not found.")
        try:
            if finds["scanduration"]:
                self.scanduration_textvar.set(finds["scanduration"].group(1))
        except Exception:
            print("Scan duration not found.")

        self.indicator_light_importdata.itemconfig(self.import_data_light, fill="green")

    def shift_log_file(self):
        self.target_dir = str(fdlg.askdirectory())
        self.shift_log_file_label.configure(text=str(self.target_dir))
        if str(self.target_dir) != "":
            self.indicator_light_shift_logfile.itemconfig(
                self.shift_log_file_light, fill="green"
            )

    def select_batch_file(self):
        filepath = fdlg.askopenfilename(
            filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")]
        )
        if not filepath:
            return
        self.batch_listbox.delete(0, tk.END)
        self.batch_file_list = []
        try:
            with open(filepath, "r") as f:
                lines = [line.strip() for line in f.readlines() if line.strip()]
            for line in lines:
                self.batch_file_list.append(line)
                self.batch_listbox.insert(tk.END, Path(line).name)
            self.select_batch_file_label.configure(text=filepath)
        except Exception as e:
            messagebox.showerror("Error", f"Could not read batch file: {e}")

    def on_batch_select(self, event):
        selection = self.batch_listbox.curselection()
        if not selection:
            return
        self.batch_selected_index = selection[0]
        filepath = self.batch_file_list[self.batch_selected_index]
        if not Path(filepath).exists():
            messagebox.showerror("Error", f"File not found:\n{filepath}")
            return
        self.selected_log_source_file = filepath
        self.select_log_file_label.configure(text=filepath)
        self.indicator_light_importdata.itemconfig(
            self.import_data_light, fill="yellow"
        )
        self.parse_data()

    def return_to_batch(self):
        self.batch_listbox.selection_clear(0, tk.END)
        self.reset_clear()

    def publish_info(self):
        self.data_dict["Proposal Number"] = self.proposalnumber_textvar.get()
        self.data_dict["Name"] = self.name_textvar.get()
        self.data_dict["Date"] = self.date_textvar.get()

        position_parts = []
        if self.facultypos_intvar.get() == 1:
            position_parts.append("Faculty")
        if self.staffpos_intvar.get() == 1:
            position_parts.append("Staff")
        if self.studentpos_intvar.get() == 1:
            position_parts.append("Student")
        if self.otherpos_intvar.get() == 1:
            position_parts.append("Other")
        self.data_dict["Position"] = "/".join(position_parts)

        self.data_dict["Operator"] = self.operator_textvar.get()
        self.data_dict["NYIT or External"] = self.nyitorext_textvar.get()
        self.data_dict["Filename"] = self.filename_textvar.get()
        self.data_dict["Voltage"] = self.voltage_textvar.get()
        self.data_dict["Current"] = self.current_textvar.get()

        temp_res = self.resolution_textvar.get()
        if temp_res not in ("0.5K", "1K", "2K"):
            res_map = {2240: "2K", 1120: "1K", 560: "0.5K"}
            try:
                act_res = res_map.get(int(temp_res), "None")
            except ValueError:
                act_res = "None"
        else:
            act_res = temp_res
        self.data_dict["Resolution"] = act_res

        self.data_dict["Filter"] = self.filter_textvar.get()
        self.data_dict["Exposure"] = self.exposure_textvar.get()
        self.data_dict["Pixel Size"] = self.pixelsize_textvar.get()
        self.data_dict["Image Format"] = self.imageformat_textvar.get()
        self.data_dict["Type of Scan"] = self.typeofscan_textvar.get()
        if self.offset_yesno.get() == 1:
            self.data_dict["Type of Scan"] += "; Offset"
        self.data_dict["Rotation Step"] = self.rotationstep_textvar.get()
        self.data_dict["Frames"] = self.frames_textvar.get()
        self.data_dict["Random Movement"] = self.randommovement_textvar.get()
        self.data_dict["Rotate 360"] = (
            "yes" if self.rotate360_yesno.get() == 1 else "no"
        )
        self.data_dict["Scan Duration"] = self.scanduration_textvar.get()
        comments = self.comments_input.get("1.0", "end-1c")
        self.data_dict["Comments"] = "" if comments == "Enter Here" else comments
        self.data_dict["Log File Archived"] = (
            "yes" if self.logfilearchive_yesno.get() == 1 else "no"
        )
        self.data_dict["Viscenter Storage"] = (
            "yes" if self.visstorage_yesno.get() == 1 else "no"
        )

        appendable_data = [
            self.data_dict["Proposal Number"],
            self.data_dict["Name"],
            self.data_dict["Date"],
            self.data_dict["Position"],
            self.data_dict["NYIT or External"],
            self.data_dict["Filename"],
            self.data_dict["Voltage"],
            self.data_dict["Current"],
            self.data_dict["Resolution"],
            self.data_dict["Filter"],
            self.data_dict["Exposure"],
            self.data_dict["Pixel Size"],
            self.data_dict["Image Format"],
            self.data_dict["Type of Scan"],
            self.data_dict["Rotation Step"],
            self.data_dict["Frames"],
            self.data_dict["Random Movement"],
            self.data_dict["Rotate 360"],
            self.data_dict["Scan Duration"],
            self.data_dict["Comments"],
            self.data_dict["Log File Archived"],
            self.data_dict["Viscenter Storage"],
            self.data_dict["Operator"],
        ]

        myxlfile = Path(self.target_archive_file)
        wb = openpyxl.load_workbook(myxlfile)
        sheet = wb[wb.sheetnames[0]]
        sheetmaxrow = sheet.max_row + 1
        for i in range(sheet.max_column):
            column_let = get_column_letter(i + 1)
            if 0 <= i < len(appendable_data):
                sheet[f"{column_let}{sheetmaxrow}"] = appendable_data[i]
        wb.save(myxlfile)

        if str(self.selected_log_source_file) != "":
            shutil.copy(str(self.selected_log_source_file), str(self.target_dir))

        if self.batch_file_list and self.batch_selected_index >= 0:
            self.batch_listbox.delete(self.batch_selected_index)
            self.batch_file_list.pop(self.batch_selected_index)
            self.batch_selected_index = -1
            self.return_to_batch()

    def reset_clear(self):
        if self.destinationfilekeep_yesno.get() != 1:
            self.select_destination_file_label.configure(text="")
            self.indicator_light_destination_file.itemconfig(
                self.destination_file_light, fill="red"
            )
        self.select_log_file_label.configure(text="")
        self.indicator_light_importdata.itemconfig(self.import_data_light, fill="red")
        if self.destinationdirectorykeep_yesno.get() != 1:
            self.shift_log_file_label.configure(text="")
            self.indicator_light_shift_logfile.itemconfig(
                self.shift_log_file_light, fill="red"
            )

        self.proposalnumber_textvar.set("Enter Here")
        self.name_textvar.set("Enter Here")
        self.date_textvar.set("Auto Fill")
        self.facultypos_intvar.set(0)
        self.staffpos_intvar.set(0)
        self.studentpos_intvar.set(0)
        self.otherpos_intvar.set(0)
        self.operator_textvar.set("Enter Here")
        self.nyitorext_textvar.set("Enter Here")
        self.filename_textvar.set("Auto Fill")
        self.voltage_textvar.set("Auto Fill")
        self.current_textvar.set("Auto Fill")
        self.resolution_textvar.set("Click to Select")
        self.filter_textvar.set("Enter or Select Here")
        self.exposure_textvar.set("Auto Fill")
        self.pixelsize_textvar.set("Auto Fill")
        self.imageformat_textvar.set("Auto Fill")
        self.typeofscan_textvar.set("Enter Here")
        self.offset_yesno.set(0)
        self.rotationstep_textvar.set("Auto Fill")
        self.frames_textvar.set("Auto Fill")
        self.randommovement_textvar.set("Auto Fill")
        self.rotate360_yesno.set(0)
        self.scanduration_textvar.set("Auto Fill")
        self.comments_input.delete(1.0, tk.END)
        self.comments_input.insert(tk.END, "Enter Here")
        self.logfilearchive_yesno.set(0)
        self.visstorage_yesno.set(0)


# ==============================================================================
# TAB 4: Log File Export Tool
# ==============================================================================

# Regex patterns for Export Tool log parsing
EXPORT_LOG_PATTERNS = {
    "DateTime": r"Date and Time\s*=\s*(.+)",
    "Filename": r"Filename Prefix\s*=\s*(.+)",
    "Voltage": r"Source Voltage\s*\(kV\)\s*=\s*(.+)",
    "Current": r"Source Current\s*\(uA\)\s*=\s*(.+)",
    "PixelSize": r"Image Pixel Size\s*\(um\)\s*=\s*(.+)",
    "Filter": r"Filter\s*=\s*(.+)",
    "Exposure": r"Exposure\s*\(ms\)\s*=\s*(.+)",
    "RotStep": r"Rotation Step\s*\(deg\)=(.+)",
    "FrameAvg": r"Frame Averaging\s*=\s*(.+)",
    "RandMov": r"Random Movement\s*=\s*(.+)",
    "Rot360": r"Use 360 Rotation\s*=\s*(.+)",
}


class ExportTool(ttk.Frame):
    """Log File Export Tool — fills a Word doc template with parsed log data."""

    def __init__(self, parent):
        super().__init__(parent)
        self.targ_dir = ""
        self.targ_file = ""
        # Declare entry widgets with type hints for linter resolution
        self.exp_date: ttk.Entry
        self.exp_filename: ttk.Entry
        self.exp_scannedby: ttk.Entry
        self.exp_machine: ttk.Entry
        self.exp_voltage: ttk.Entry
        self.exp_current: ttk.Entry
        self.exp_pixelsize: ttk.Entry
        self.exp_filter: ttk.Entry
        self.exp_exposure: ttk.Entry
        self.exp_rotationstep: ttk.Entry
        self.exp_frameaverage: ttk.Entry
        self.exp_randmovement: ttk.Entry
        self.rb_var: tk.StringVar
        self.exp_addcomment: tk.Text
        self._build_ui()

    def _build_ui(self):
        # Title
        title_frame = ttk.Frame(master=self)
        title_frame.grid(column=0, row=0, columnspan=6, sticky=tk.W + tk.E)
        title_frame.columnconfigure(0, weight=1)
        ttk.Label(
            master=title_frame,
            text="Log File Export Tool",
            relief=tk.FLAT,
            font=("Times New Roman", 18, "bold"),
        ).grid(column=0, row=0, columnspan=6, pady=(10, 0))

        ttk.Separator(master=self, orient="horizontal").grid(
            row=1, column=0, columnspan=6, sticky="nesw", pady=(5, 5)
        )

        # Top buttons row
        btn_frame = ttk.Frame(master=self)
        btn_frame.grid(
            row=2, column=0, columnspan=6, padx=5, pady=5, sticky=tk.W + tk.E
        )
        btn_frame.columnconfigure(2, weight=1)

        ttk.Button(
            master=btn_frame, text="Load Log File", command=self.load_log_file
        ).grid(column=0, row=0, padx=5, pady=5, sticky=tk.W)

        ttk.Button(
            master=btn_frame,
            text="Select Target Directory",
            command=self.select_destination,
        ).grid(column=1, row=0, padx=5, pady=5, sticky=tk.W)

        self.target_dir_label_frame = ttk.LabelFrame(
            master=btn_frame, text="Target Directory"
        )
        self.target_dir_label_frame.grid(
            column=2, row=0, columnspan=3, padx=5, pady=5, sticky=tk.W + tk.E
        )
        self.target_dir_label = ttk.Label(
            master=self.target_dir_label_frame, text="", width=80
        )
        self.target_dir_label.grid(row=0, column=0, padx=5, pady=2, sticky=tk.W)

        ttk.Separator(master=self, orient="horizontal").grid(
            row=3, column=0, columnspan=6, sticky="nesw", pady=(0, 5)
        )

        # Input fields arranged in a 2-column grid
        input_frame = ttk.Frame(master=self)
        input_frame.grid(
            row=4, column=0, columnspan=6, padx=10, pady=5, sticky=tk.W + tk.E
        )

        field_defs = [
            ("Date:", "exp_date"),
            ("File Name(s):", "exp_filename"),
            ("Scanned By:", "exp_scannedby"),
            ("Machine:", "exp_machine"),
            ("Voltage (kV):", "exp_voltage"),
            ("Current (uA):", "exp_current"),
            ("Pixel Size (\u03bcm):", "exp_pixelsize"),
            ("Filter:", "exp_filter"),
            ("Exposure (ms):", "exp_exposure"),
            ("Rotation Step (deg):", "exp_rotationstep"),
            ("Frame Averaging:", "exp_frameaverage"),
            ("Random Movement:", "exp_randmovement"),
        ]

        self.entries = {}
        for i, (label_text, attr_name) in enumerate(field_defs):
            row = i // 2
            col_base = (i % 2) * 3
            ttk.Label(master=input_frame, text=label_text).grid(
                row=row, column=col_base, padx=5, pady=3, sticky=tk.E
            )
            entry = ttk.Entry(master=input_frame, width=30)
            entry.grid(
                row=row, column=col_base + 1, columnspan=2, padx=5, pady=3, sticky=tk.W
            )
            self.entries[attr_name] = entry
            setattr(self, attr_name, entry)

        # 360 Rotation radio buttons
        rot_row = len(field_defs) // 2
        ttk.Label(master=input_frame, text="360\u00b0 Rotation?:").grid(
            row=rot_row, column=0, padx=5, pady=3, sticky=tk.E
        )
        self.rb_var = tk.StringVar(value="Yes")
        ttk.Radiobutton(
            master=input_frame, text="Yes", variable=self.rb_var, value="Yes"
        ).grid(row=rot_row, column=1, padx=5, pady=3, sticky=tk.W)
        ttk.Radiobutton(
            master=input_frame, text="No", variable=self.rb_var, value="No"
        ).grid(row=rot_row, column=2, padx=5, pady=3, sticky=tk.W)

        # Additional Comments
        ttk.Label(
            master=self,
            text="Additional Comments:",
            font=("Times New Roman", 9, "bold underline"),
        ).grid(row=5, column=0, columnspan=2, padx=10, pady=(5, 0), sticky=tk.W)
        self.exp_addcomment = tk.Text(master=self, height=8, width=70)
        self.exp_addcomment.grid(
            row=6, column=0, columnspan=4, padx=10, pady=(0, 10), sticky=tk.W
        )

        ttk.Separator(master=self, orient="horizontal").grid(
            row=7, column=0, columnspan=6, sticky="nesw", pady=(0, 5)
        )

        # Bottom action buttons
        action_frame = ttk.Frame(master=self)
        action_frame.grid(row=8, column=0, columnspan=6, padx=10, pady=5, sticky=tk.E)

        ttk.Button(
            master=action_frame,
            text="Select Template File",
            command=self.load_template_file,
        ).grid(row=0, column=0, padx=5, pady=5)

        ttk.Button(
            master=action_frame, text="Publish", command=self.publish_log_file
        ).grid(row=0, column=1, padx=5, pady=5)

        ttk.Button(master=action_frame, text="Reset", command=self.reset_fields).grid(
            row=0, column=2, padx=5, pady=5
        )

    def load_log_file(self):
        filepath = fdlg.askopenfilename(
            filetypes=[
                ("Log Files", "*.log"),
                ("Text Files", "*.txt"),
                ("All Files", "*.*"),
            ]
        )
        if not filepath:
            return

        with open(filepath, "r") as f:
            text = f.read()

        results = {
            key: re.search(pattern, text)
            for key, pattern in EXPORT_LOG_PATTERNS.items()
        }

        if not all(results.values()):
            missing = [k for k, v in results.items() if not v]
            messagebox.showerror(
                "Parse Error", f"Could not find fields: {', '.join(missing)}"
            )
            return

        data = {key: match.group(1).strip() for key, match in results.items() if match}

        self.reset_fields()

        self.exp_date.insert(0, data["DateTime"])
        self.exp_filename.insert(0, data["Filename"])
        self.exp_scannedby.insert(0, "Dan Gibbons")
        self.exp_machine.insert(0, "Bruker Skyscan 1173")
        self.exp_voltage.insert(0, data["Voltage"])
        self.exp_current.insert(0, data["Current"])
        self.exp_pixelsize.insert(0, data["PixelSize"])
        self.exp_filter.insert(0, data["Filter"])
        self.exp_exposure.insert(0, data["Exposure"])
        self.exp_rotationstep.insert(0, data["RotStep"])
        self.exp_frameaverage.insert(0, data["FrameAvg"])
        self.exp_randmovement.insert(0, data["RandMov"])
        self.rb_var.set("Yes" if data["Rot360"].upper() == "YES" else "No")

    def select_destination(self):
        directory = fdlg.askdirectory()
        if directory:
            self.targ_dir = directory
            self.target_dir_label.configure(text=directory)

    def load_template_file(self):
        filepath = fdlg.askopenfilename(
            filetypes=[("Word Doc Files", "*.docx"), ("All Files", "*.*")]
        )
        if not filepath:
            return
        if not self.targ_dir:
            messagebox.showerror("Error", "Please select a target directory first.")
            return
        self.targ_file = self.targ_dir + "/External_User_Info.docx"
        copyfile(filepath, self.targ_file)
        messagebox.showinfo("Done", f"Template copied to:\n{self.targ_file}")

    def publish_log_file(self):
        if not self.targ_file:
            messagebox.showerror("Error", "No template file selected.")
            return

        values = {
            label: widget.get()
            for label, widget in [
                ("Date: ", self.exp_date),
                ("File Name(s): ", self.exp_filename),
                ("Scanned by: ", self.exp_scannedby),
                ("Machine: ", self.exp_machine),
                ("Voltage (kV): ", self.exp_voltage),
                ("Current (uA): ", self.exp_current),
                ("Pixel Size (\u03bcm): ", self.exp_pixelsize),
                ("Filter: ", self.exp_filter),
                ("Exposure (ms): ", self.exp_exposure),
                ("Rotation Step (deg): ", self.exp_rotationstep),
                ("Frame Averaging: ", self.exp_frameaverage),
                ("Random Movement: ", self.exp_randmovement),
            ]
        }
        values["360\u00b0 Rotation?: "] = self.rb_var.get()
        values["Additional Comments: "] = self.exp_addcomment.get("1.0", tk.END).strip()

        try:
            doc = Document(self.targ_file)

            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(11)

            for paragraph in doc.paragraphs:
                for label, value in values.items():
                    if label in paragraph.text:
                        paragraph.text = paragraph.text.replace(label, label + value)
                        paragraph.style = "Normal"
                        break

            doc.save(self.targ_file)
            messagebox.showinfo("Success", "Document published successfully!")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to publish document:\n{e}")

    def reset_fields(self):
        for entry in self.entries.values():
            entry.delete(0, tk.END)
        self.exp_addcomment.delete("1.0", tk.END)
        self.rb_var.set("Yes")


# ==============================================================================
# Main Application Window
# ==============================================================================


class App(tk.Tk):
    """Main application window containing all three tools as notebook tabs."""

    def __init__(self):
        super().__init__()
        self.title("MicroCT Log File Tools")
        self._build_ui()

    def _build_ui(self):
        # App-level title
        ttk.Label(
            master=self,
            text="MicroCT Log File Tools",
            font=("Times New Roman", 20, "bold"),
        ).grid(row=0, column=0, pady=(10, 0))

        ttk.Label(
            master=self,
            text="Designed and Created by Dan Gibbons  |  Property of NYIT",
            font=("Times New Roman", 9, "italic"),
        ).grid(row=1, column=0, pady=(0, 5))

        ttk.Separator(master=self, orient="horizontal").grid(
            row=2, column=0, sticky="nesw", pady=(0, 5)
        )

        # Notebook
        notebook = ttk.Notebook(master=self)
        notebook.grid(row=3, column=0, sticky=tk.W + tk.E + tk.N + tk.S, padx=5, pady=5)

        tab_extract = ExtractTool(notebook)
        tab_check = CheckTool(notebook)
        tab_parse = ParseTool(notebook)
        tab_export = ExportTool(notebook)

        notebook.add(tab_extract, text="  Log File Extractor  ")
        notebook.add(tab_check, text="  Log File Checker  ")
        notebook.add(tab_parse, text="  Log File Parser  ")
        notebook.add(tab_export, text="  Log File Exporter  ")


if __name__ == "__main__":
    app = App()
    app.mainloop()
